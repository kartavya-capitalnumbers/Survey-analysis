"""
stats_report.py — narrative + statistical DOCX report for the survey-data app.

Per the spec's "Narrative summary" output: a plain-English analytical report
structured by topic, in a tone suitable for direct inclusion in an ESIA, ESDD,
or lender submission. This is the statistical-app counterpart to report.py
(which serves the narrative Mode-2 app): it combines everything already
computed elsewhere in this app —
    demographics.run_breakdowns()   -> demographic tables
    qualitative.tag_responses() +
        summarise_themes()          -> ranked themes, sentiment, quotes
    charts.py                       -> bar/pie/pyramid PNGs, embedded inline
into one DOCX. The executive summary reuses stats.write_narrative_summary() —
the same grounded, multi-paragraph RAP/ESIA-style narrative generator the
"Ask in plain English" tab uses for narrative_summary questions, so both
surfaces produce writing of the same depth and quality from the same figures.
"""

from __future__ import annotations

import io
from typing import List, Optional

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt

import charts
import demographics
import qualitative
import stats


# ---------------------------------------------------------------------------
# DOCX assembly helpers
# ---------------------------------------------------------------------------

def _add_df_table(document: Document, df: pd.DataFrame) -> None:
    if df.empty:
        document.add_paragraph("(no data)")
        return
    table = document.add_table(rows=1, cols=len(df.columns))
    table.style = "Light Grid Accent 1"
    for i, col in enumerate(df.columns):
        cell = table.rows[0].cells[i]
        cell.text = str(col)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            cells[i].text = "" if pd.isna(row[col]) else str(row[col])


def _add_stat_result(document: Document, result: stats.StatResult) -> None:
    p = document.add_paragraph()
    p.add_run(result.title).bold = True
    _add_df_table(document, result.table)
    cap = document.add_paragraph()
    cap.add_run(f"n = {result.n} · {result.methodology}").italic = True
    for c in result.caveats:
        cw = document.add_paragraph()
        cw.add_run(f"Caveat: {c}").italic = True


def _add_png(document: Document, png_bytes: bytes, *, width_in: float = 5.5) -> None:
    document.add_picture(io.BytesIO(png_bytes), width=Inches(width_in))


# ---------------------------------------------------------------------------
# Orchestration
# ---------------------------------------------------------------------------

def make_stats_report_docx(
    dataset_name: str,
    df: pd.DataFrame,
    members_df: Optional[pd.DataFrame] = None,
) -> bytes:
    """Build the full narrative + statistical DOCX report and return its bytes.

    Runs the demographic breakdowns and thematic/sentiment synthesis itself
    (via demographics.py / qualitative.py) so callers only need to supply the
    dataset — everything downstream is grounded in what those modules compute.
    """
    demo_results = demographics.run_breakdowns(df)

    tags: List[qualitative.ResponseTag] = []
    theme_summaries: List[qualitative.ThemeSummary] = []
    if "project_concern" in df.columns:
        tags = qualitative.tag_responses(df)
        theme_summaries = qualitative.summarise_themes(tags)

    # Pass the already-computed theme_summaries so write_narrative_summary
    # doesn't re-run qualitative classification a second time.
    exec_summary = stats.write_narrative_summary(df, theme_summaries=theme_summaries)

    document = Document()
    document.add_heading("Household Survey — Statistical & Qualitative Analysis", level=0)
    subtitle = document.add_paragraph()
    run = subtitle.add_run(f"Dataset: {dataset_name} · {len(df)} households")
    run.italic = True
    run.font.size = Pt(11)

    # --- 1. Executive summary ------------------------------------------------
    document.add_heading("1. Executive summary", level=1)
    for para_text in exec_summary.split("\n\n"):
        if para_text.strip():
            document.add_paragraph(para_text.strip())

    # --- 2. Demographic breakdowns --------------------------------------------
    document.add_heading("2. Demographic breakdowns", level=1)
    document.add_paragraph(
        "Gender, age group, vulnerability category, and geographic zone, per the "
        "standard breakdown set."
    )
    for i, result in enumerate(demo_results, start=1):
        document.add_heading(f"2.{i} {result.title}", level=2)
        _add_stat_result(document, result)

    # --- 3. Charts -------------------------------------------------------------
    if "village" in df.columns:
        document.add_heading("3. Charts", level=1)
        try:
            freq = stats.frequency(df, "village")
            _add_png(document, charts.chart_to_png(
                charts.bar_chart(freq.table, "village", "count", title="Households by village")
            ))
        except stats.StatsError:
            pass
        if "vulnerable_household" in df.columns:
            try:
                freq2 = stats.frequency(df, "vulnerable_household")
                _add_png(document, charts.chart_to_png(
                    charts.pie_chart(freq2.table, "vulnerable_household", "count",
                                      title="Vulnerable households")
                ))
            except stats.StatsError:
                pass
        if members_df is not None and not members_df.empty:
            _add_png(document, charts.chart_to_png(charts.demographic_pyramid(members_df)))

    # --- 4. Qualitative synthesis ----------------------------------------------
    if theme_summaries:
        document.add_heading("4. Qualitative synthesis", level=1)
        document.add_paragraph(
            f"{len(tags)} free-text responses classified into "
            f"{len(theme_summaries)} themes. Quotes below are anonymised."
        )
        document.add_heading("4.1 Ranked concern themes", level=2)
        _add_df_table(document, qualitative.themes_table(theme_summaries))

        document.add_heading("4.2 Sentiment breakdown", level=2)
        _add_df_table(document, qualitative.sentiment_table(tags))

        document.add_heading("4.3 Representative quotes", level=2)
        for s in theme_summaries:
            document.add_paragraph(f"{s.theme} ({s.count}, {s.pct}%)", style="List Bullet")
            for q in s.sample_quotes:
                qp = document.add_paragraph(f"“{q}”", style="List Bullet 2")

    # --- 5. Appendix -------------------------------------------------------------
    document.add_heading("5. Methodology & limitations", level=1)
    document.add_paragraph(
        "All statistics in this report are computed deterministically from the "
        "extracted dataset; the language model is used only to plan/narrate "
        "already-computed figures and to classify/anonymise free-text responses "
        "— it never derives a number itself. Percentages are of valid "
        "(non-missing) responses for each field. Tables with small per-row "
        "sample sizes should be treated as indicative — see the confidence "
        "flags in the Excel export. The full de-identified dataset and, where "
        "explicitly exported, the raw transcription are available separately."
    )

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()
